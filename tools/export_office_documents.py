from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


ROOT = Path(__file__).resolve().parents[1]
TEST_BOOK_SOURCE = ROOT / "SYSTEM_TEST_CASE_BOOK.md"
DOCUMENTATION_SOURCE = ROOT / "SYSTEM_DOCUMENTATION_NOTE.md"
TEST_BOOK_OUTPUT = ROOT / "SYSTEM_TEST_CASE_BOOK.xlsx"
DOCUMENTATION_OUTPUT = ROOT / "SYSTEM_DOCUMENTATION_NOTE.docx"

HEADER_FILL = PatternFill("solid", fgColor="1F4E78")
SECTION_FILL = PatternFill("solid", fgColor="D9EAF7")
PASS_FILL = PatternFill("solid", fgColor="C6EFCE")
FAIL_FILL = PatternFill("solid", fgColor="FFC7CE")
BLOCKED_FILL = PatternFill("solid", fgColor="FFEB9C")


def clean_cell(value):
    value = value.strip()
    value = re.sub(r"`([^`]+)`", r"\1", value)
    return value.replace("\\|", "|")


def parse_markdown_table(lines):
    rows = []
    for line in lines:
        if not line.startswith("|"):
            continue
        cells = [clean_cell(cell) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def parse_test_book(source):
    sections = []
    current_section = None
    current_table = []

    for line in source.read_text(encoding="utf-8").splitlines():
        if line.startswith("## "):
            if current_section:
                current_section["tables"].append(parse_markdown_table(current_table))
                sections.append(current_section)
            current_section = {"title": line[3:].strip(), "tables": []}
            current_table = []
        elif line.startswith("|") and current_section:
            current_table.append(line)
        elif current_table and current_section:
            current_section["tables"].append(parse_markdown_table(current_table))
            current_table = []

    if current_section:
        current_section["tables"].append(parse_markdown_table(current_table))
        sections.append(current_section)

    return sections


def style_header(row):
    for cell in row:
        cell.fill = HEADER_FILL
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def size_columns(worksheet):
    for column_index, column_cells in enumerate(worksheet.columns, start=1):
        column_letter = get_column_letter(column_index)
        longest = max(len(str(cell.value or "")) for cell in column_cells)
        worksheet.column_dimensions[column_letter].width = min(max(longest + 2, 14), 55)


def add_execution_sheet(workbook, sections):
    worksheet = workbook.active
    worksheet.title = "Execution Summary"
    worksheet.append(["MediCore SaaS Test Execution"])
    worksheet.merge_cells("A1:D1")
    worksheet["A1"].fill = HEADER_FILL
    worksheet["A1"].font = Font(color="FFFFFF", bold=True, size=14)
    worksheet["A1"].alignment = Alignment(horizontal="center")
    worksheet.append([])
    worksheet.append(["Field", "Value", "Field", "Value"])
    style_header(worksheet[3])
    fields = [
        ("Build/version", "", "Environment", ""),
        ("Tester", "", "Date", ""),
        ("Hospital under test", "", "Browser/device", ""),
        ("Overall result", "", "Defect references", ""),
    ]
    for row in fields:
        worksheet.append(row)
    worksheet.append([])
    worksheet.append(["Section", "Test cases", "Passed", "Failed"])
    style_header(worksheet[9])
    for section in sections:
        test_count = sum(max(len(table) - 1, 0) for table in section["tables"] if table)
        worksheet.append([section["title"], test_count, "", ""])
    size_columns(worksheet)
    worksheet.freeze_panes = "A4"


def add_test_cases_sheet(workbook, sections):
    worksheet = workbook.create_sheet("Test Cases")
    worksheet.append([
        "Section",
        "Test ID",
        "Scenario",
        "Steps",
        "Expected Result",
        "Status",
        "Tester Notes",
        "Defect ID",
    ])
    style_header(worksheet[1])
    status_validation = DataValidation(
        type="list",
        formula1='"Not Run,Pass,Fail,Blocked,Not Applicable"',
        allow_blank=True,
    )
    worksheet.add_data_validation(status_validation)

    for section in sections:
        for table in section["tables"]:
            if not table or table[0][0] != "ID":
                continue
            for row in table[1:]:
                values = row + [""] * (4 - len(row))
                worksheet.append([
                    section["title"],
                    values[0],
                    values[1],
                    values[2],
                    values[3],
                    "Not Run",
                    "",
                    "",
                ])

    final_row = worksheet.max_row
    status_validation.add(f"F2:F{final_row}")
    worksheet.conditional_formatting.add(
        f"F2:F{final_row}",
        FormulaRule(formula=['F2="Pass"'], fill=PASS_FILL),
    )
    worksheet.conditional_formatting.add(
        f"F2:F{final_row}",
        FormulaRule(formula=['F2="Fail"'], fill=FAIL_FILL),
    )
    worksheet.conditional_formatting.add(
        f"F2:F{final_row}",
        FormulaRule(formula=['F2="Blocked"'], fill=BLOCKED_FILL),
    )
    worksheet.auto_filter.ref = f"A1:H{final_row}"
    worksheet.freeze_panes = "A2"
    worksheet.sheet_view.showGridLines = False
    size_columns(worksheet)
    worksheet.column_dimensions["A"].width = 32
    worksheet.column_dimensions["B"].width = 14
    worksheet.column_dimensions["C"].width = 28
    worksheet.column_dimensions["D"].width = 52
    worksheet.column_dimensions["E"].width = 60
    worksheet.column_dimensions["F"].width = 18
    worksheet.column_dimensions["G"].width = 35
    worksheet.column_dimensions["H"].width = 16
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def add_reference_sheet(workbook, title, rows):
    worksheet = workbook.create_sheet(title)
    for row in rows:
        worksheet.append(row)
    if rows:
        style_header(worksheet[1])
    worksheet.freeze_panes = "A2"
    worksheet.sheet_view.showGridLines = False
    size_columns(worksheet)
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def create_test_workbook():
    sections = parse_test_book(TEST_BOOK_SOURCE)
    workbook = Workbook()
    add_execution_sheet(workbook, sections)
    add_test_cases_sheet(workbook, sections)

    for section in sections:
        for table in section["tables"]:
            if table and table[0][0] == "Field":
                add_reference_sheet(workbook, "Test Setup", table)
            elif table and table[0][0] == "Account":
                add_reference_sheet(workbook, "Test Accounts", table)

    workbook.save(TEST_BOOK_OUTPUT)


def add_word_table(document, rows):
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    document.add_paragraph()


def create_documentation_word_file():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    lines = DOCUMENTATION_SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code_block = False
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue
        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            index += 1
            continue
        if line.startswith("# "):
            paragraph = document.add_heading(clean_cell(line[2:]), level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            document.add_heading(clean_cell(line[3:]), level=1)
        elif line.startswith("### "):
            document.add_heading(clean_cell(line[4:]), level=2)
        elif line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = parse_markdown_table(table_lines)
            if rows:
                add_word_table(document, rows)
            continue
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(clean_cell(re.sub(r"^\d+\. ", "", line)), style="List Number")
        elif line.startswith("- "):
            document.add_paragraph(clean_cell(line[2:]), style="List Bullet")
        elif line.strip():
            document.add_paragraph(clean_cell(line))
        index += 1

    document.save(DOCUMENTATION_OUTPUT)


if __name__ == "__main__":
    create_test_workbook()
    create_documentation_word_file()
    print(f"Created {TEST_BOOK_OUTPUT.name}")
    print(f"Created {DOCUMENTATION_OUTPUT.name}")